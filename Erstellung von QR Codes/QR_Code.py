#%%
import qrcode
import os
from PIL import Image, ImageDraw, ImageFont
import pandas as pd
from docx import Document
from docx.shared import Cm
#%%
def Data_to_QR(Messpunkte,parent_dir,StrName):
    try:
        os.mkdir(parent_dir+"\\QR-Codes-PNG-"+StrName)
    except:
        print("Ordner: QR-Codes-PNG ist vorhanden")
    for i in range(1,Messpunkte.size+1):
        #creating a variable for the text in which we want the QR code to link us
        data=Messpunkte[i]
        qr = qrcode.make(data)

        qr_Temp = Image.new("RGB", (qr.size[0],qr.size[1]), (255, 255, 255))
        qr_Temp.paste(qr)

        # Creating single QRImage

        qr=qr_Temp.resize((250,250))

        #fnt = ImageFont.truetype("FONTS\\arial\\arial", 30)
        fnt = ImageFont.truetype("FONTS\\arial\\arial", 25)
        d = ImageDraw.Draw(qr)

        d.multiline_text((125, 245), data, font=fnt, fill=(0, 0, 0),align="center",anchor="ms")
        d.multiline_text((125, 25), data, font=fnt, fill=(0, 0, 0),align="center",anchor="ms")

        qr=qr.transpose(Image.Transpose.ROTATE_270)
        onlyqr= Image.new("RGB", (286,199), (255, 255, 255))
        onlyqr.paste(qr,(30,-25))

        # Creating QRImage for Avery

        out = Image.new("RGB", (573,200), (255, 255, 255))

        out.paste(onlyqr,(-15,0))

        onlyqr = onlyqr.transpose(Image.Transpose.ROTATE_180)
        out.paste(onlyqr,(301,0))

        d = ImageDraw.Draw(out)
        d.line((287, 0,287,201), fill=0,width=5)
        #d.rectangle((-1,-1,out.size[0],out.size[1]),fill=None,outline=55,width=5)

        file_name=parent_dir+"\\QR-Codes-PNG-"+StrName+"\\"+data+".png"
        out.save(file_name)

    return

# %%
def Print_to_Pdf(Messpunkte): #Depricated

    i=2
    fak_inch_to_mm=0.0394
    Pixel_x=[595,1240,2480]
    Pixel_y=[842,1754,3508]
    DPI=[72,150,300]
    fak=fak_inch_to_mm*DPI[i]

    try:
        os.mkdir(parent_dir+"\\QR-Codes-PDF")
    except:
        print("Ordner: QR-Codes-PNG ist vorhanden")

    out_end=Image.new("RGB", (Pixel_x[i], Pixel_y[i]), (255, 255, 255)) 
    #box=(9*fak,13.5*fak,(Pixel_x[i]-6*fak),(Pixel_y[i]-13*fak)) für den Drucker Zuhause
    box=((9-1.5)*fak,(13.5-2.0)*fak,(Pixel_x[i]-6*fak),(Pixel_y[i]-13*fak)) #für den Drucker auf der Arbeit
    #9,13.5,-6,-13

    #d = ImageDraw.Draw(out_end)
    #d.rectangle(box,fill=None,outline=128,width=1)

    #out_end.save(r"C:\Users\Richard Brand\HiDrive\Pyhton\Python Projekts\QR-CODE\PDF Barcodes\Seitetest.pdf",resolution=DPI[i])

    iter_a=0
    while iter_a<Messpunkte.size:
        for k in range(0,4):
            for n in range(0,16):
                iter_a+=1
                
                x=int(k*573)
                y=int(n*200)
                try:
                    out=Image.open(str(parent_dir+"\\QR-Codes-PNG\\"+Messpunkte[iter_a]+".png"))
                except:
                    print("Fertich")


                out_end.paste(out,(x+int(box[0]),y+int(box[1])))
                
                b=iter_a%64
                b_n=iter_a%Messpunkte.size
                #if iter_a==Messpunkte.size:
                    #return                
                #print(iter_a)
                if  iter_a==64 or b==0 or iter_a==Messpunkte.size: #Messpunkte.size
                    d = ImageDraw.Draw(out_end)
                    d.rectangle(box,fill=None,outline=128,width=5)
                    #out_end.show()
                    out_end.save(str(parent_dir+"\\QR-Codes-PDF\\"+"Seite "+str(iter_a-63)+" - "+str(iter_a)+".pdf"),resolution=DPI[i])
                    out_end=Image.new("RGB", (Pixel_x[i], Pixel_y[i]), (255, 255, 255))
                    if iter_a==Messpunkte.size:
                        return


#%%
def PngToWord(Messpunkte,StrParentdir,StrName):
    #Parentdir=r"C:\Users\Richard.STB-HALLE\HiDrive\Pyhton\Python Projekts\QR-CODE\Fertige Etiketten\Brücken Etiketten\QR-Codes-PNG\BR-W1-A-1.png"
    # Liste mit Bildpfaden
    left_margin = Cm(0.8)
    right_margin = Cm(0.8)
    top_margin = Cm(1.15)
    bottom_margin = Cm(1.45)

    image_paths =[]
    for x in range(1,len(Messpunkte)+1):
        StrPathTemp=str(StrParentdir+"\\QR-Codes-PNG-"+StrName+"\\"+Messpunkte[x]+".png")
        image_paths.append(StrPathTemp)

    # Öffnen des vorhandenen Word-Dokuments
    IntZahl=1
    count=0
    while len(image_paths)>0:
        document = Document(str(StrParentdir+"\\TEMPLATES\\AveryZweckform3667_TemplateNew.docx"))
        # Iterieren über die Zellen der Tabelle
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Bildpfad für die aktuelle Zelle
                    if len(image_paths)==0:
                        break

                    image_path = image_paths.pop(0)
                    count+=1
                    #count=+
                    #count+=1

                    # Bild in die Zelle einfügen
                    cell_paragraph = cell.paragraphs[0]
                    cell_run = cell_paragraph.add_run()
                    cell_run.add_picture(image_path, width=Cm(4.72), height=Cm(1.65))

                    if count%64==0:
                        break
        
        FileName="{} Etiketten Teil {}".format(StrName,IntZahl)
        StrEndName="{}\\{}.docx".format(StrParentdir,FileName)
        sections = document.sections
        for section in sections:
            section.left_margin = left_margin
            section.right_margin = right_margin
            section.top_margin = top_margin
            section.bottom_margin = bottom_margin
        document.save(StrEndName)
        IntZahl+=1
        if len(image_paths)==0:
            break

#%%

# URL für wo sich diese Datei befindet
parent_dir=r"U:\User\_HIWIS\Sarah\Erstellung von QR Codes"
# URL wo sich die CSV-befindet
path=r"U:\User\_HIWIS\Sarah\Erstellung von QR Codes\Modell1_Coburg_W1 Abschnitt 1.csv"
#C:\Users\Richard.STB-HALLE\HiDrive\Pyhton\Python Projekts\QR-CODE\QR-Codes-PNG
StrName="W1 Abschnitt 1"

B=pd.read_csv(path,names=["Messpunkt"])
Messpunkte=B.iloc[1:,0]

#%%
Data_to_QR(Messpunkte,parent_dir,StrName)

# %%
#Print_to_Pdf(Messpunkte=Messpunkte)
PngToWord(Messpunkte,parent_dir,StrName)
# %%

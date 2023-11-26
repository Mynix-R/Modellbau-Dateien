from PIL import Image
from pybarcodes import EAN13
from PIL import Image, ImageDraw, ImageFont
import os
from docx import Document
from docx.shared import Cm
import pandas as pd


def NumberProof(barcodenumber):

    barcode = str(barcodenumber)  # Der Barcode ohne Prüfziffer

    # Schritt 1: Zahlen an ungeraden Positionen addieren
    odd_sum = sum(int(barcode[i]) for i in range(0, 12, 2))

    # Schritt 2: Zahlen an geraden Positionen addieren
    even_sum = sum(int(barcode[i]) for i in range(1, 12, 2))

    # Schritt 3: Summe der Schritte 1 und 2 multiplizieren
    total_sum = odd_sum + even_sum*3
    # Schritt 4: Nächstgelegene Zahl finden, die ein Vielfaches von 10 ist
    nearest_ten = (total_sum // 10 + 1) * 10

    # Schritt 5: Prüfziffer berechnen
    check_digit = nearest_ten - total_sum

    if check_digit==10:
        check_digit=0

    return str(check_digit)


def Data_to_Barcode(NameCSV,StrName,BaseNumberFile):
    DirFilePath = os.getcwd()
    CSVPath=DirFilePath+"\\"+NameCSV
    Messpunkte=pd.read_csv(CSVPath,names=["Messpunkt"])
    Messpunkte["IdentNumber"]=0

    try:
        os.mkdir(DirFilePath+"\\BQ-Codes-PNG-"+StrName)
    except:
        print("Ordner: QR-Codes-PNG ist vorhanden")

    TXtPath=DirFilePath+"\\RESSOURCES\\"+BaseNumberFile

    with open(TXtPath) as f:
        Line = f.readlines()
    f.close
    if len(Line[0])!=12:
        print(f"IntBasenumber has {len(Line)} not 12 as expected")
    IntBasenumber=[int(x) for x in Line][0]
    IntAdd=10

    for i in range(1,Messpunkte.shape[0]):
        IntBasenumber+=IntAdd
        
        #creating a variable for the text in which we want the QR code to link us
        CODE=str(IntBasenumber)+NumberProof(IntBasenumber)
        Messpunkte.iloc[i,1]=int(CODE)
        Messpunkt=Messpunkte.iloc[i,0] 
        barcode = EAN13(CODE)

        #qr_Temp = Image.new("RGB", (barcode.size[0],barcode.size[1]), (255, 255, 255))
        #qr_Temp.paste(barcode)
        barcode.BARCODE_FONT_SIZE=0
        barcode.BARCODE_COLUMN_NUMBER= 110
        #barcode.BARCODE_FONT_SIZE= 46
        barcode.BARCODE_LENGTH= 12
        #barcode.BARCODE_PADDING= (100, 200)
        barcode.BARCODE_SIZE= (720, 360)
        #barcode.FIRST_SECTION= (0, 6)
        barcode.HAS_STRUCTURE= True
        #barcode.SECOND_SECTION= (6, 12)
        barcode.WEIGHTS= (3, 1)

        file_name=DirFilePath+"\\BQ-Codes-PNG-"+StrName+"\\"+Messpunkt+".png"
        barcode.save(path=file_name)
        # PNG-Bild öffnen
        image = Image.open(file_name)
        image=image.resize((250,250))

        fnt = ImageFont.truetype("FONTS\\arial\\arial", 25)
        d = ImageDraw.Draw(image)

        d.multiline_text((125, 245), Messpunkt, font=fnt, fill=(0, 0, 0),align="center",anchor="ms")
        d.multiline_text((125, 25), str(CODE), font=fnt, fill=(0, 0, 0),align="center",anchor="ms")

        qr=image.transpose(Image.Transpose.ROTATE_270)
        onlyqr= Image.new("RGB", (286,199), (255, 255, 255))
        onlyqr.paste(qr,(30,-25))
        # Creating QRImage for Avery

        out = Image.new("RGB", (573,200), (255, 255, 255))

        out.paste(onlyqr,(-15,0))

        onlyqr = onlyqr.transpose(Image.Transpose.ROTATE_180)
        out.paste(onlyqr,(301,0))

        d = ImageDraw.Draw(out)
        d.line((287, 0,287,201), fill=0,width=5)
        # Verarbeite das Bild (z.B. zeige es an oder führe andere Operationen durch)
        #out.show()
    # Optional: Speichere das Bild als neues PNG
        out.save(file_name)

    #Erstellen eines Ordners:
    file_name_dir=DirFilePath+"\\CompleteData-"+StrName
    DirName="\\CompleteData-{}".format(StrName)
    try:
        os.mkdir(file_name_dir)
    except:
        print("Ordner: {}".format(DirName))

    file_name_dirCSV=file_name_dir+"\\"+StrName+"-IdentNumber.csv"
    
    Messpunkte.iloc[1:,:].to_csv(file_name_dirCSV,index=False,header=False,sep=';')

    PngToWord(NameCSV,StrName)



    with open(TXtPath, 'w') as f:
            f.write(str(IntBasenumber))
    f.close
    
    return print("Das projekt {} ist Abgeschlossen".format(StrName))

def PngToWord(NameCSV,StrName):
    StrParentdir = os.getcwd()
    #Parentdir=r"C:\Users\Richard.STB-HALLE\HiDrive\Pyhton\Python Projekts\QR-CODE\Fertige Etiketten\Brücken Etiketten\QR-Codes-PNG\BR-W1-A-1.png"
    CSVPath=StrParentdir+"\\"+NameCSV
    Messpunkte=pd.read_csv(CSVPath,names=["Messpunkt"])
    Messpunkte=Messpunkte.iloc[1:,0]

    # Liste mit Bildpfaden
    left_margin = Cm(0.8)
    right_margin = Cm(0.8)
    top_margin = Cm(1.15)
    bottom_margin = Cm(1.45)

    image_paths =[]
    for x in range(1,len(Messpunkte)+1):
        StrPathTemp=str(StrParentdir+"\\BQ-Codes-PNG-"+StrName+"\\"+Messpunkte[x]+".png")
        image_paths.append(StrPathTemp)

    # Öffnen des vorhandenen Word-Dokuments
    IntZahl=1
    count=0
    while len(image_paths)>0:
        document = Document(str(StrParentdir+"\\TEMPLATES\\AveryZweckform3667_TemplateNew.docx"))
        # Iterieren über die Zellen der Tabelle
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Bildpfad für die aktuelle Zelle
                    if len(image_paths)==0:
                        break

                    image_path = image_paths.pop(0)
                    count+=1
                    #count=+
                    #count+=1

                    # Bild in die Zelle einfügen
                    cell_paragraph = cell.paragraphs[0]
                    cell_run = cell_paragraph.add_run()
                    cell_run.add_picture(image_path, width=Cm(4.72), height=Cm(1.65))

                    if count%64==0:
                        break
        
        file_name_dir=StrParentdir+"\\CompleteData-"+StrName
        FileName="{} Etiketten Teil {}".format(StrName,IntZahl)
        StrEndName="{}\\{}.docx".format(file_name_dir,FileName)
        sections = document.sections
        for section in sections:
            section.left_margin = left_margin
            section.right_margin = right_margin
            section.top_margin = top_margin
            section.bottom_margin = bottom_margin
        document.save(StrEndName)
        IntZahl+=1
        if len(image_paths)==0:
            break

    #Löschen der PNG-Barcodes und deren zugegöriger Ordner
    PathFolder=StrParentdir+"\\BQ-Codes-PNG-"+StrName
    items = os.listdir(PathFolder)

    for x in items:
        FileName=str(PathFolder)+"\\"+x
        os.remove(FileName)

    os.rmdir(StrParentdir+"\\BQ-Codes-PNG-"+StrName)
    
    
if __name__ == "__main__":
    # Die Code-Datei muss im selben Ordner sein wie die CSV!

    #Name der Unique-Basenumber:
    BaseNumber='IntBasenumber.txt'
    #Name der CSV-Datei:
    NameCSV="TestCSV.csv"
    #Name des Projekts:
    ProjektName="W1 Abschnitt 1"


    Data_to_Barcode(NameCSV,ProjektName,BaseNumber)
